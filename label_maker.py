import os
import customtkinter as ctk
from tkinter import filedialog, messagebox, colorchooser
import logging

# Set up logging to log errors to a file
logging.basicConfig(filename='file_processing_errors.log',
                    level=logging.ERROR,
                    format='%(asctime)s - %(levelname)s - %(message)s')

# Function to locate resource files, works for both PyInstaller executable and dev environment
def resource_path(relative_path):
    import sys
    try:
        # PyInstaller creates a temporary folder and stores resources there
        base_path = sys._MEIPASS
    except AttributeError:
        # If not running in PyInstaller, use the regular path
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# Ensure the history file is stored in a user-writable location (not inside the executable)
def get_history_file_path():
    """ Get the writable path for the history file in a fixed directory on your computer. """
    # Define the fixed path where you want the history file to be saved
    base_path = r"G:\Shared drives\Scribe Workspace\Scribe Master Folder\Scribe Label Maker"
        
    # Ensure the directory exists (creates it if it doesn't exist)
    os.makedirs(base_path, exist_ok=True)
    
    # Construct the full path for the history file
    history_file_path = os.path.join(base_path, "order_history.json")
    
    return history_file_path

# Set CustomTkinter appearance mode and color theme
ctk.set_appearance_mode("light")
ctk.set_default_color_theme("green")

# Global variables
labels_data = []
displayed_envelope_files = set()
displayed_letter_files = set()
order_colors = {}

# Paths for docx files
blank_template_path = resource_path('resources/Label_Template_BLANK.docx')
generated_template_path = resource_path('resources/GENERATED_Label_Template.docx')

# Function to load order history from JSON file
def load_order_history():
    import json
    history_file_path = get_history_file_path()
    if os.path.exists(history_file_path):
        with open(history_file_path, 'r') as file:
            return json.load(file)
    return []

# Function to save order history to JSON file
def save_order_history(history):
    import json
    history_file_path = get_history_file_path()
    with open(history_file_path, 'w') as file:
        json.dump(history, file)

# Function to update the order history with the last 10 entries
def update_order_history(order_name, color):
    history = load_order_history()

    # Remove any existing entries for the same order_name to ensure uniqueness
    history = [entry for entry in history if entry['order_name'] != order_name]

    # Add the new entry
    history.append({"order_name": order_name, "color": color})

    # Ensure only the last 20 unique entries are kept
    if len(history) > 20:
        history = history[-20:]

    save_order_history(history)

# Function to display the last 10 order_name and color combinations in the GUI
def display_order_history():
    # Clear previous displayed history
    for widget in history_label_frame.winfo_children():
        widget.destroy()

    history = load_order_history()

    # Reverse the history to show newest entries first
    history.reverse()

    for entry in history:
        order_name = entry['order_name']
        color = entry['color']

        # Populate the order_colors dictionary
        order_colors[order_name] = color

        # Determine the appropriate text color based on background brightness
        def is_light_color(hex_color):
            hex_color = hex_color.lstrip("#")
            r, g, b = (
                int(hex_color[0:2], 16),
                int(hex_color[2:4], 16),
                int(hex_color[4:6], 16),
            )
            brightness = (r * 299 + g * 587 + b * 114) / 1000  # Luminance formula
            return brightness > 186

        text_color = "black" if is_light_color(color) else "white"

        # Create a label for each history entry with the background color set to the assigned color
        order_label = ctk.CTkLabel(
            history_label_frame,
            text=order_name,
            font=("Helvetica", 12),
            text_color=text_color,
        )
        order_label.configure(fg_color=f"#{color}")  # Set the background color to the assigned color
        order_label.pack(pady=2, padx=5, anchor="w", fill="x")

        # Adjust bindtags to include history_label_frame
        order_label.bindtags((str(order_label), str(history_label_frame), "all"))

        # Bind click event to change color
        order_label.bind(
            "<Button-1>",
            lambda event, name=order_name, label=order_label: change_order_history_color(event, name, label),
        )

# Function to create styled text for each field in the DOCX file
def create_styled_text(order_name, batch_chip, card_envelope):
    from docxtpl import RichText
    font_color = order_colors.get(order_name, "000000")
    styled_order_name = RichText(order_name, color=font_color, size=32)
    styled_batch_chip = RichText(batch_chip, color=font_color, size=32)
    styled_card_envelope = RichText(card_envelope, color=font_color, size=32)

    return {
        'order_name': styled_order_name,
        'batch_chip': styled_batch_chip,
        'card_envelope': styled_card_envelope,
    }

# Function to select Envelope Chip files
def select_envelope_files():
    files = filedialog.askopenfilenames(title="Select Envelope Chip Files")
    if files:
        generate_labels_data(files, "Envelopes")
    else:
        messagebox.showerror("Error", "No Envelope Chip files selected!")

# Function to select Letter Chip files
def select_letter_files():
    files = filedialog.askopenfilenames(title="Select Letter Chip Files")
    if files:
        generate_labels_data(files, "Letters")
    else:
        messagebox.showerror("Error", "No Letter Chip files selected!")

# Function to generate the labels data based on selected files
def generate_labels_data(files, chip_type):
    import logging
    from collections import defaultdict
    global labels_data, displayed_envelope_files, displayed_letter_files, order_colors
    valid_files = []  # To store valid files for display in the labels
    order_type_count = defaultdict(list)  # To group by order_name and card_envelope type
    invalid_files = []  # List to track invalid files

    for file_path in files:
        try:
            # Validate if file exists and is accessible
            if not os.path.isfile(file_path):
                logging.error(f"File not found or inaccessible: {file_path}")
                continue

            file_name = os.path.basename(file_path)
            base_name = os.path.splitext(file_name)[0]  # Remove the extension (e.g., ".bin")
            
            # Ignore hyphens in base_name
            base_name = base_name.replace("-", " ")

            # Ensure the files match the correct chip_type (Envelopes or Letters)
            if chip_type == "Envelopes" and "Letters" in base_name:
                invalid_files.append(file_name)  # Track invalid files
                continue  # Skip this file as it's not valid for Envelopes
            elif chip_type == "Letters" and "Envelopes" in base_name:
                invalid_files.append(file_name)  # Track invalid files
                continue  # Skip this file as it's not valid for Letters

            # Splitting the file name into parts
            parts = base_name.split()

            # Extract the order_name and card_envelope type
            if "Envelopes" in base_name:
                order_name = base_name.split("Envelopes")[0].strip()  # Get everything before "Envelopes"
                card_envelope = "Envelope"
            elif "Letters" in base_name:
                order_name = base_name.split("Letters")[0].strip()  # Get everything before "Letters"
                card_envelope = "Card"

            # The rest of the logic remains unchanged
            order_type_count[(order_name, card_envelope)].append(file_name)

            # Prompt for color if order_name doesn't already have one
            if order_name not in order_colors:
                assign_color_for_order(order_name)

        except Exception as e:
            # Log any unexpected exceptions during file processing
            logging.error(f"Failed to process file '{file_path}': {str(e)}")

    # Notify the user if there were any invalid files
    if invalid_files:
        invalid_file_list = "\n".join(invalid_files)
        messagebox.showerror("Invalid Files", f"The following files were invalid for {chip_type}:\n{invalid_file_list}")

    # Assign batch numbers for each unique order_name and type
    for (order_name, card_envelope), file_list in order_type_count.items():
        for i, file_name in enumerate(file_list):
            batch_chip = f"{i + 1} of {len(file_list)}"

            # Append valid files to labels_data
            labels_data.append({
                "order_name": order_name,
                "batch_chip": batch_chip,
                "card_envelope": card_envelope
            })

            # Add valid files to the appropriate set for preventing duplicates
            if chip_type == "Envelopes" and file_name not in displayed_envelope_files:
                valid_files.append(file_name)
                displayed_envelope_files.add(file_name)
            elif chip_type == "Letters" and file_name not in displayed_letter_files:
                valid_files.append(file_name)
                displayed_letter_files.add(file_name)

    # Display only valid and non-duplicate files in the appropriate label
    if valid_files:
        file_names = "\n".join(valid_files)
        if chip_type == "Envelopes":
            current_text = envelope_label.cget("text")
            envelope_label.configure(text=current_text + file_names + "\n")
        elif chip_type == "Letters":
            current_text = letter_label.cget("text")
            letter_label.configure(text=current_text + file_names + "\n")

# Function to prompt user for a color for each unique order_name
def assign_color_for_order(order_name):
    color = colorchooser.askcolor(title=f"Choose color for {order_name}")
    if color[1]:
        order_colors[order_name] = color[1][1:]
        display_order_color(order_name, color[1])

# Function to change the color when the label is clicked
def change_color(order_name, color_label):
    color = colorchooser.askcolor(title=f"Choose a new color for {order_name}")
    if color[1]:
        order_colors[order_name] = color[1][1:]
        color_label.configure(fg_color=color[1])

# Function to display the order color without right-click options
def display_order_color(order_name, color_hex):
    def is_light_color(hex_color):
        hex_color = hex_color.lstrip("#")
        r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        brightness = (r * 299 + g * 587 + b * 114) / 1000  # Luminance formula
        return brightness > 186

    text_color = "black" if is_light_color(color_hex) else "white"

    envelope_label.pack_forget()
    letter_label.pack_forget()

    color_label = ctk.CTkLabel(
        scrollable_frame,
        text=f"{order_name} assigned color",
        fg_color=color_hex,
        text_color=text_color
    )
    color_label.pack(pady=5, padx=20)

    letter_label.pack(pady=10, padx=20, fill="x", side="bottom")
    envelope_label.pack(pady=10, padx=20, fill="x", side="bottom")

def change_order_history_color(event, order_name, order_label):
    color = colorchooser.askcolor(title=f"Choose new color for {order_name}")
    if color[1]:
        new_color = color[1][1:]  # Remove the '#' from the color code
        order_colors[order_name] = new_color
        order_label.configure(fg_color=color[1])

        # Update text color based on brightness
        def is_light_color(hex_color):
            hex_color = hex_color.lstrip("#")
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            brightness = (r * 299 + g * 587 + b * 114) / 1000
            return brightness > 186

        text_color = "black" if is_light_color(new_color) else "white"
        order_label.configure(text_color=text_color)

        # Update the order history
        update_order_history(order_name, new_color)

# Function to create labels in the DOCX file with font size 16 and bold formatting (excluding curly braces)
def create_labels(doc, start_number, num_labels):
    for i in range(start_number, start_number + num_labels):
        order_paragraph = add_label_paragraph(doc, f"Order Name & Number\n{{{{ order_name{i} }}}}")
        batch_chip_paragraph = add_label_paragraph(doc, f"Batch / Chip Number\n{{{{ batch_chip{i} }}}}")
        type_paragraph = add_label_paragraph(doc, f"Type: {{{{ card_envelope{i} }}}}")

# Function to add paragraphs and apply bold formatting
def add_label_paragraph(doc, text):
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    paragraph = doc.add_paragraph()
    is_bold = True
    current_run = ""
    
    for idx, char in enumerate(text):
        if char == '{' and idx+1 < len(text) and text[idx+1] == '{':
            if current_run:
                run = paragraph.add_run(current_run)
                run.bold = is_bold
                run.font.size = Pt(16)
            current_run = char
            is_bold = False
        elif char == '}' and idx-1 >= 0 and text[idx-1] == '}':
            current_run += char
            run = paragraph.add_run(current_run)
            run.bold = is_bold
            run.font.size = Pt(16)
            current_run = ""
            is_bold = True
        else:
            current_run += char

    if current_run:
        run = paragraph.add_run(current_run)
        run.bold = is_bold
        run.font.size = Pt(16)
    
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return paragraph

# Function to create and save the DOCX file and display the "Open File" button if successful
def create_docx():
    if not labels_data:
        messagebox.showerror("Error", "No valid files selected!")
        return

    try:
        from docx import Document
        from docxtpl import DocxTemplate, RichText
        import logging
        doc = Document(blank_template_path)
        if len(doc.paragraphs) > 0:
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        num_labels_to_generate = len(labels_data)
        create_labels(doc, 1, num_labels_to_generate)
        doc.save(generated_template_path)

        template = DocxTemplate(generated_template_path)
        context = {}
        for i, label in enumerate(labels_data):
            context[f'order_name{i+1}'] = RichText(label['order_name'], color=order_colors[label['order_name']], size=32, bold=True)
            context[f'batch_chip{i+1}'] = RichText(label['batch_chip'], color=order_colors[label['order_name']], size=32, bold=True)
            context[f'card_envelope{i+1}'] = RichText(label['card_envelope'], color=order_colors[label['order_name']], size=32, bold=True)

            update_order_history(label['order_name'], order_colors[label['order_name']])

        template.render(context)
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if save_path:
            template.save(save_path)
            messagebox.showinfo("Success", f"Labels saved to {save_path}")
            open_button.configure(command=lambda: open_docx_file(save_path))
            open_button.pack(pady=10, padx=20, before=reset_button)
        else:
            messagebox.showerror("Error", "Save path not specified or operation cancelled.")

        display_order_history()

    except Exception as e:
        logging.error(f"Failed to create the DOCX file: {str(e)}")
        messagebox.showerror("Error", f"Failed to create the DOCX file: {str(e)}")

# Function to open the DOCX file
def open_docx_file(file_path):
    import platform
    import logging
    try:
        if platform.system() == "Windows":
            os.startfile(file_path)
        elif platform.system() == "Darwin":
            os.system(f"open {file_path}")
        else:
            os.system(f"xdg-open {file_path}")
    except Exception as e:
        logging.error(f"Failed to open the DOCX file: {str(e)}")
        messagebox.showerror("Error", f"Failed to open the DOCX file: {str(e)}")

# Scroll event for mouse (Windows)
def on_mousewheel(event):
    canvas.yview_scroll(int(-1*(event.delta/120)), "units")

# Scroll event for macOS (Mac uses different event types for scrolling)
def on_mousewheel_mac(event):
    canvas.yview_scroll(-1 if event.num == 5 else 1, "units")

# Function to reset label data, displayed files, and order colors
def reset_data():
    global labels_data, displayed_envelope_files, displayed_letter_files, order_colors

    # Clear all relevant data
    labels_data.clear()
    displayed_envelope_files.clear()
    displayed_letter_files.clear()
    order_colors.clear()

    # Reset label texts
    envelope_label.configure(text="Selected Envelope Files:\n")
    letter_label.configure(text="Selected Letter Files:\n")

    # Destroy any color assignment labels
    for widget in scrollable_frame.winfo_children():
        if isinstance(widget, ctk.CTkLabel) and "assigned color" in widget.cget("text"):
            widget.destroy()

    # Hide the open button
    open_button.pack_forget()

    messagebox.showinfo("Reset", "All label data and displayed files have been reset!")

# GUI Setup
root = ctk.CTk()
root.title("Label Maker")
root.geometry("700x600")  # Increased the window width to accommodate both sides
root.configure(bg="#3A3A3A")

icon_path = resource_path('resources/scribe-icon.ico')
root.iconbitmap(icon_path)

logo_path = resource_path('resources/scribe-logo-final.png')

from PIL import Image
logo_image = Image.open(logo_path)
logo_image = logo_image.resize((258, 100), Image.Resampling.LANCZOS)

logo_ctk_image = ctk.CTkImage(light_image=logo_image, dark_image=logo_image, size=(258, 100))
logo_label = ctk.CTkLabel(root, image=logo_ctk_image, text="")
logo_label.pack(pady=10)

# Left side scrollable frame for main content
left_frame = ctk.CTkFrame(root, fg_color="#3A3A3A", corner_radius=15)
left_frame.pack(side="left", fill="both", expand=True, padx=(10, 5), pady=10)

# Create an inner frame to hold the canvas and leave space for the corner radius
inner_frame = ctk.CTkFrame(left_frame, fg_color="#3A3A3A", corner_radius=15)
inner_frame.pack(expand=True, fill="both", padx=10, pady=10)

canvas = ctk.CTkCanvas(inner_frame, bg="#3A3A3A", highlightthickness=0)
scrollbar = ctk.CTkScrollbar(inner_frame, orientation="vertical", command=canvas.yview)
scrollable_frame = ctk.CTkFrame(canvas, fg_color="#3A3A3A", corner_radius=15)

scrollable_frame.bind(
    "<Configure>",
    lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
)

canvas.create_window((0, 0), window=scrollable_frame, anchor="nw", width=450)
canvas.configure(yscrollcommand=scrollbar.set)

canvas.bind_all("<MouseWheel>", on_mousewheel)
canvas.bind_all("<Button-4>", on_mousewheel_mac)
canvas.bind_all("<Button-5>", on_mousewheel_mac)

canvas.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")

# Right side frame for order history
right_frame = ctk.CTkFrame(root, fg_color="#2E2E2E", corner_radius=15, width=250)
right_frame.pack(side="right", fill="y", padx=(5, 10), pady=10)

history_label_frame = ctk.CTkScrollableFrame(right_frame, width=230, fg_color="#2E2E2E", height=500)
history_label_frame.pack(pady=10, padx=10, fill="both", expand=True)

# Add widgets to left scrollable frame
instruction = ctk.CTkLabel(scrollable_frame, text="Select files to generate labels", font=("Helvetica", 16), text_color="white")
instruction.pack(pady=1, padx=20, expand=False)

# Add a new label for the file format below the instruction label
file_format_example = ctk.CTkLabel(scrollable_frame, text="Examples: Chris LaVigne T1 Copy A Envelopes-1-167.bin\n" + 
                                                                "Chris LaVigne T1 List 1 Letters-1-167.bin\n" +
                                                                    "Chris LaVigne T1 Envelopes-1-167.bin", 
                                 font=("Helvetica", 12), text_color="gray")
file_format_example.pack(pady=1, padx=20, expand=False)

envelope_button = ctk.CTkButton(scrollable_frame, text="Select Envelope Chip Files", command=select_envelope_files)
envelope_button.pack(pady=10, padx=20, fill="x", expand=True)

letter_button = ctk.CTkButton(scrollable_frame, text="Select Letter Chip Files", command=select_letter_files)
letter_button.pack(pady=10, padx=20, fill="x", expand=True)

create_button = ctk.CTkButton(scrollable_frame, text="Create DOCX", command=create_docx, fg_color="#133d8e", hover_color="#266cc3")
create_button.pack(pady=10, padx=20, fill="x", expand=True)

reset_button = ctk.CTkButton(scrollable_frame, text="Reset", command=reset_data, width=100, fg_color="#8e1313", hover_color="#c32626")
reset_button.pack(pady=15, padx=20)

open_button = ctk.CTkButton(scrollable_frame, text="Open Created DOCX File", width=300, fg_color="#133d8e", hover_color="#266cc3")
open_button.pack_forget()

envelope_label = ctk.CTkLabel(scrollable_frame, text="Selected Envelope Files:\n", font=("Helvetica", 12), text_color="white", anchor="w", justify="left")
envelope_label.pack(pady=10, padx=20, fill="x", side="top")

letter_label = ctk.CTkLabel(scrollable_frame, text="Selected Letter Files:\n", font=("Helvetica", 12), text_color="white", anchor="w", justify="left")
letter_label.pack(pady=10, padx=20, fill="x", side="top")

# Display the order history on the right side
display_order_history()

root.mainloop()